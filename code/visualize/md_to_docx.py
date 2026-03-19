"""
Convert 投稿主稿_v1.md to a formatted Word document.
Follows 《计算机工程与应用》 submission style:
  A4, 2.5cm margins, SimSun body, 1.5 line spacing.

Usage:
    python code/visualize/md_to_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
MD_PATH = ROOT / "docs" / "投稿主稿_v1.md"
OUT_PATH = ROOT / "docs" / "投稿主稿_v1.docx"
FIGURES_DIR = ROOT / "figures"


def set_page_layout(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_default_font(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)


def setup_heading_styles(doc):
    for level, sz in [(1, 16), (2, 14), (3, 13)]:
        name = f"Heading {level}"
        style = doc.styles[name]
        font = style.font
        font.name = "黑体"
        font.size = Pt(sz)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        pf = style.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5


def add_title_block(doc, title, author, affiliation):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(author)
    run2.font.size = Pt(12)
    run2.font.name = "宋体"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(affiliation)
    run3.font.size = Pt(10)
    run3.font.name = "宋体"
    run3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run3.font.color.rgb = RGBColor(100, 100, 100)
    p3.paragraph_format.space_after = Pt(16)


def add_abstract_block(doc, abstract_text, keywords):
    p = doc.add_paragraph()
    run_label = p.add_run("摘要：")
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = "黑体"
    run_label._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run_body = p.add_run(abstract_text)
    run_body.font.size = Pt(12)
    run_body.font.name = "宋体"
    run_body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.first_line_indent = Cm(0.74)

    pk = doc.add_paragraph()
    run_kl = pk.add_run("关键词：")
    run_kl.bold = True
    run_kl.font.size = Pt(12)
    run_kl.font.name = "黑体"
    run_kl._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run_kv = pk.add_run(keywords)
    run_kv.font.size = Pt(12)
    run_kv.font.name = "宋体"
    run_kv._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pk.paragraph_format.space_after = Pt(12)


def add_english_abstract(doc, abstract_text, keywords):
    p = doc.add_paragraph()
    run_label = p.add_run("Abstract: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.name = "Times New Roman"
    run_body = p.add_run(abstract_text)
    run_body.font.size = Pt(11)
    run_body.font.name = "Times New Roman"
    p.paragraph_format.first_line_indent = Cm(0.74)

    pk = doc.add_paragraph()
    run_kl = pk.add_run("Keywords: ")
    run_kl.bold = True
    run_kl.font.size = Pt(11)
    run_kl.font.name = "Times New Roman"
    run_kv = pk.add_run(keywords)
    run_kv.font.size = Pt(11)
    run_kv.font.name = "Times New Roman"
    pk.paragraph_format.space_after = Pt(16)


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            clean = re.sub(r"\$[^$]+\$", lambda m: m.group(0).strip("$"), clean)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            run = p.add_run(clean)
            run.font.size = Pt(10)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def add_figure_placeholder(doc, fig_num, caption):
    fig_file = None
    fig_map = {
        "1": "fig1_runtime_comparison_v3.png",
        "2": "fig2_turncount_comparison_v3.png",
        "3": "fig3_ablation_study_v3.png",
        "4": "path_comparison_6maps_v3.png",
        "5": "fig5_nodes_vs_obstacle_v3.png",
    }
    if fig_num in fig_map:
        candidate = FIGURES_DIR / fig_map[fig_num]
        if candidate.exists():
            fig_file = candidate

    if fig_file:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig_file), width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[图 {fig_num} 占位：请在此插入图片]")
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.font.size = Pt(11)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = cap.add_run(f"图 {fig_num}  {caption}")
    run_cap.font.size = Pt(10)
    run_cap.font.name = "宋体"
    run_cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.paragraph_format.space_after = Pt(8)


def clean_inline_formatting(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def main():
    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()
    set_page_layout(doc)
    set_default_font(doc)
    setup_heading_styles(doc)

    add_title_block(
        doc,
        "面向移动机器人的自适应加权 A* 路径规划算法研究",
        "刘佳淑",
        "（作者单位，城市 邮编）",
    )

    cn_abstract = ""
    cn_keywords = ""
    en_abstract = ""
    en_keywords = ""
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line == "## 摘要":
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("**关键词"):
                if lines[i].strip() and not lines[i].strip().startswith(">") and not lines[i].strip() == "---":
                    buf.append(lines[i].strip())
                i += 1
            cn_abstract = " ".join(buf)
            if i < len(lines) and "关键词" in lines[i]:
                cn_keywords = re.sub(r"\*\*关键词[：:]\*\*\s*", "", lines[i].strip())
            break
        i += 1

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line == "## Abstract":
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("**Keywords"):
                if lines[i].strip() and not lines[i].strip() == "---":
                    buf.append(lines[i].strip())
                i += 1
            en_abstract = " ".join(buf)
            if i < len(lines) and "Keywords" in lines[i]:
                en_keywords = re.sub(r"\*\*Keywords[：:]\*\*\s*", "", lines[i].strip())
            break
        i += 1

    add_abstract_block(doc, clean_inline_formatting(cn_abstract), clean_inline_formatting(cn_keywords))
    add_english_abstract(doc, clean_inline_formatting(en_abstract), clean_inline_formatting(en_keywords))

    skip_sections = {"摘要", "Abstract", "图表清单", "复现命令"}
    in_skip = False
    in_code_block = False
    in_table = False
    table_buf = []
    figure_inserted = set()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            in_code_block = True
            if "text" in stripped or not stripped[3:].strip():
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph()
                for cl in code_lines:
                    run = p.add_run(cl + "\n")
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1.0)
                in_code_block = False
                if i < len(lines):
                    i += 1
                continue
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        if stripped == "---" or stripped.startswith("> 数据口径"):
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            if heading_text in skip_sections:
                in_skip = True
                i += 1
                continue
            in_skip = False

        if in_skip:
            i += 1
            continue

        if stripped.startswith("|") and not in_table:
            in_table = True
            table_buf = [stripped]
            i += 1
            continue
        if in_table:
            if stripped.startswith("|"):
                table_buf.append(stripped)
                i += 1
                continue
            else:
                rows = parse_table(table_buf)
                add_table_to_doc(doc, rows)
                in_table = False
                table_buf = []

        fig_match = re.search(r"[（(]见图\s*(\d+)[）)]|图\s*(\d+)\s", stripped)
        if fig_match and not stripped.startswith("##"):
            fig_num = fig_match.group(1) or fig_match.group(2)
            if fig_num and fig_num not in figure_inserted:
                pass

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            heading_text = re.sub(r"^第\d+章\s*", "", heading_text)
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            heading_text = stripped[5:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        if stripped.startswith("$$"):
            formula = stripped.strip("$ ")
            if not formula:
                i += 1
                formula_lines = []
                while i < len(lines) and not lines[i].strip().startswith("$$"):
                    formula_lines.append(lines[i].strip())
                    i += 1
                formula = " ".join(formula_lines)
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.italic = True
            i += 1 if i < len(lines) else i
            continue

        if stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s*", "", stripped)
            run = p.add_run(clean_inline_formatting(text))
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            text = stripped[2:]
            run = p.add_run(clean_inline_formatting(text))
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            i += 1
            continue

        if stripped and not stripped.startswith("["):
            text = clean_inline_formatting(stripped)
            text = re.sub(r"\$([^$]+)\$", r"\1", text)

            fig_ref = re.search(r"[（(]见图\s*(\d+)[）)]", text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            p.paragraph_format.first_line_indent = Cm(0.74)

            if fig_ref:
                fn = fig_ref.group(1)
                if fn not in figure_inserted:
                    captions = {
                        "1": "各算法运行时间对比（按地图类型）",
                        "2": "各算法转弯次数对比（按地图类型）",
                        "3": "消融实验结果（全局15张地图均值）",
                        "4": "路径对比可视化（传统 A* vs 改进 A*）",
                        "5": "扩展节点数随障碍率的变化趋势",
                    }
                    add_figure_placeholder(doc, fn, captions.get(fn, ""))
                    figure_inserted.add(fn)

        i += 1

    if in_table and table_buf:
        rows = parse_table(table_buf)
        add_table_to_doc(doc, rows)

    doc.save(str(OUT_PATH))
    print(f"[OK] Word document saved: {OUT_PATH}")
    print(f"     Size: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
